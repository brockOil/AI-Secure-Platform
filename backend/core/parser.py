from fastapi import UploadFile
import io

SUPPORTED_FILE_TYPES = {".pdf", ".docx", ".txt", ".log"}

async def parse_upload(file: UploadFile) -> tuple[str, str]:
    filename = file.filename or ""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    raw_bytes = await file.read()

    content = ""
    # Try parsing based on extension
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=raw_bytes, filetype="pdf")
            for page in doc:
                text = page.get_text()
                if text:
                    content += text + "\n"
        except Exception as e:
            content = f"Error reading PDF: {str(e)}"
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            for para in doc.paragraphs:
                content += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            content += para.text + "\n"
        except Exception as e:
            content = f"Error reading DOCX: {str(e)}"
    else:
        # Fallback to UTF-8 decoding for txt/log or unknown files
        content = raw_bytes.decode("utf-8", errors="replace")

    if ext in (".log", ".txt"):
        return content, "log" if ext == ".log" else "file"

    return content, "file"

def parse_text(content: str, input_type: str) -> str:
    return content.strip()
